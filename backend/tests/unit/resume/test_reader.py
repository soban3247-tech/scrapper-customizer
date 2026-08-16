from io import BytesIO
from pathlib import Path

import pymupdf as fitz
import pytest
from docx import Document

from job_assistant.resume import (
    EmptyResumeTextError,
    EncryptedResumeError,
    InvalidResumeFileError,
    ResumeFileTooLargeError,
    ResumeFormat,
    UnsupportedResumeTypeError,
    read_resume,
    read_resume_bytes,
)
from job_assistant.resume import reader


def pdf_bytes(*pages: str) -> bytes:
    document = fitz.open()
    for text in pages:
        page = document.new_page()
        if text:
            page.insert_text((72, 72), text)
    data = document.tobytes()
    document.close()
    return data


def encrypted_pdf_bytes() -> bytes:
    document = fitz.open()
    document.new_page().insert_text((72, 72), "Private CV")
    stream = BytesIO()
    document.save(
        stream,
        encryption=fitz.PDF_ENCRYPT_AES_256,
        owner_pw="owner-secret",
        user_pw="user-secret",
    )
    document.close()
    return stream.getvalue()


def docx_bytes() -> bytes:
    document = Document()
    document.sections[0].header.paragraphs[0].text = "Malik Soban Rabbani"
    document.add_heading("Experience", level=1)
    document.add_paragraph("Python Developer")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "Python, SQL"
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_reads_pdf_pages_and_normalizes_text() -> None:
    result = read_resume_bytes(
        pdf_bytes("Malik   Soban", "Python Developer"),
        "candidate.PDF",
    )

    assert result.format is ResumeFormat.PDF
    assert result.filename == "candidate.PDF"
    assert result.page_count == 2
    assert "Malik Soban" in result.text
    assert "Python Developer" in result.text


def test_reads_docx_headers_paragraphs_and_tables() -> None:
    result = read_resume_bytes(docx_bytes(), "candidate.docx")

    assert result.format is ResumeFormat.DOCX
    assert result.page_count is None
    assert "Malik Soban Rabbani" in result.text
    assert "Experience" in result.text
    assert "Python Developer" in result.text
    assert "Skills | Python, SQL" in result.text


def test_reads_resume_from_local_path(tmp_path: Path) -> None:
    path = tmp_path / "candidate.pdf"
    path.write_bytes(pdf_bytes("Backend Engineer"))

    result = read_resume(path)

    assert result.filename == "candidate.pdf"
    assert result.text == "Backend Engineer"


@pytest.mark.parametrize("filename", ["candidate.txt", "candidate.doc"])
def test_rejects_unsupported_extensions(filename: str) -> None:
    with pytest.raises(UnsupportedResumeTypeError, match="PDF and DOCX"):
        read_resume_bytes(b"plain text", filename)


@pytest.mark.parametrize(
    ("data", "filename"),
    [(b"not a PDF", "candidate.pdf"), (b"not a DOCX", "candidate.docx")],
)
def test_rejects_extension_content_mismatch(data: bytes, filename: str) -> None:
    with pytest.raises(InvalidResumeFileError):
        read_resume_bytes(data, filename)


def test_rejects_oversized_upload_before_parsing() -> None:
    with pytest.raises(ResumeFileTooLargeError, match="size limit"):
        read_resume_bytes(b"%PDF-" + b"x" * 100, "candidate.pdf", max_file_size=10)


def test_rejects_docx_whose_expanded_archive_is_too_large(monkeypatch) -> None:
    monkeypatch.setattr(reader, "MAX_DOCX_UNCOMPRESSED_SIZE", 100)

    with pytest.raises(InvalidResumeFileError, match="expanded DOCX"):
        read_resume_bytes(docx_bytes(), "candidate.docx")


def test_rejects_password_protected_pdf() -> None:
    with pytest.raises(EncryptedResumeError, match="Password-protected"):
        read_resume_bytes(encrypted_pdf_bytes(), "candidate.pdf")


def test_rejects_pdf_without_selectable_text_with_ocr_guidance() -> None:
    with pytest.raises(EmptyResumeTextError, match="OCR"):
        read_resume_bytes(pdf_bytes(""), "scanned.pdf")


def test_rejects_docx_without_readable_text() -> None:
    document = Document()
    stream = BytesIO()
    document.save(stream)

    with pytest.raises(EmptyResumeTextError, match="DOCX"):
        read_resume_bytes(stream.getvalue(), "empty.docx")


def test_rejects_missing_local_file(tmp_path: Path) -> None:
    with pytest.raises(InvalidResumeFileError, match="could not be opened"):
        read_resume(tmp_path / "missing.pdf")
