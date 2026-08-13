"""Safe text extraction from PDF and DOCX CV uploads."""

import re
import zipfile
from dataclasses import dataclass
from enum import StrEnum
from io import BytesIO
from pathlib import Path
from typing import Any

import pymupdf as fitz
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from .errors import (
    EmptyResumeTextError,
    EncryptedResumeError,
    InvalidResumeFileError,
    ResumeFileTooLargeError,
    UnsupportedResumeTypeError,
)

DEFAULT_MAX_FILE_SIZE = 10 * 1024 * 1024
MAX_DOCX_UNCOMPRESSED_SIZE = 50 * 1024 * 1024
MAX_DOCX_ARCHIVE_MEMBERS = 1_000
SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


class ResumeFormat(StrEnum):
    """CV file types supported by the MVP."""

    PDF = "pdf"
    DOCX = "docx"


@dataclass(frozen=True, slots=True)
class ResumeDocument:
    """Extracted CV text and non-sensitive document metadata."""

    filename: str
    format: ResumeFormat
    text: str
    page_count: int | None = None


def read_resume(
    path: str | Path,
    *,
    max_file_size: int = DEFAULT_MAX_FILE_SIZE,
) -> ResumeDocument:
    """Read a local PDF or DOCX CV after validating its file size and type."""

    file_path = Path(path)
    try:
        size = file_path.stat().st_size
    except OSError as exc:
        raise InvalidResumeFileError(f"CV file could not be opened: {file_path}") from exc
    _validate_size(size, max_file_size)
    try:
        data = file_path.read_bytes()
    except OSError as exc:
        raise InvalidResumeFileError(f"CV file could not be read: {file_path}") from exc
    return read_resume_bytes(data, file_path.name, max_file_size=max_file_size)


def read_resume_bytes(
    data: bytes,
    filename: str,
    *,
    max_file_size: int = DEFAULT_MAX_FILE_SIZE,
) -> ResumeDocument:
    """Read uploaded CV bytes without first writing private data to disk."""

    if not isinstance(data, bytes):
        raise TypeError("data must be bytes")
    _validate_size(len(data), max_file_size)
    suffix = Path(filename).suffix.casefold()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise UnsupportedResumeTypeError("Only PDF and DOCX CV files are supported")

    if suffix == ".pdf":
        return _read_pdf(data, filename)
    return _read_docx(data, filename)


def _read_pdf(data: bytes, filename: str) -> ResumeDocument:
    if not data.startswith(b"%PDF-"):
        raise InvalidResumeFileError("The uploaded .pdf file is not a valid PDF")
    try:
        document = fitz.open(stream=data, filetype="pdf")
    except (fitz.FileDataError, RuntimeError, ValueError) as exc:
        raise InvalidResumeFileError("The PDF is damaged or cannot be opened") from exc

    try:
        if document.needs_pass:
            raise EncryptedResumeError(
                "Password-protected PDFs are not supported; upload an unlocked copy"
            )
        page_text = [page.get_text("text") for page in document]
        text = _normalize_text("\n\n".join(page_text))
        page_count = document.page_count
    except EncryptedResumeError:
        raise
    except (RuntimeError, ValueError) as exc:
        raise InvalidResumeFileError("Text could not be extracted from the PDF") from exc
    finally:
        document.close()

    _require_text(text, ResumeFormat.PDF)
    return ResumeDocument(
        filename=Path(filename).name,
        format=ResumeFormat.PDF,
        text=text,
        page_count=page_count,
    )


def _read_docx(data: bytes, filename: str) -> ResumeDocument:
    _validate_docx_archive(data)
    try:
        document = Document(BytesIO(data))
    except (PackageNotFoundError, ValueError, KeyError, zipfile.BadZipFile) as exc:
        raise InvalidResumeFileError("The DOCX is damaged or cannot be opened") from exc

    blocks: list[str] = []
    visited_parts: set[str] = set()
    for section in document.sections:
        for container in (section.header, section.footer):
            part_name = str(container.part.partname)
            if part_name not in visited_parts:
                visited_parts.add(part_name)
                _append_docx_container(container, blocks)
    _append_docx_container(document, blocks)

    text = _normalize_text("\n".join(blocks))
    _require_text(text, ResumeFormat.DOCX)
    return ResumeDocument(
        filename=Path(filename).name,
        format=ResumeFormat.DOCX,
        text=text,
    )


def _validate_docx_archive(data: bytes) -> None:
    if not data.startswith(b"PK"):
        raise InvalidResumeFileError("The uploaded .docx file is not a valid DOCX")
    try:
        with zipfile.ZipFile(BytesIO(data)) as archive:
            members = archive.infolist()
            if len(members) > MAX_DOCX_ARCHIVE_MEMBERS:
                raise InvalidResumeFileError("The DOCX contains too many archive entries")
            if sum(member.file_size for member in members) > MAX_DOCX_UNCOMPRESSED_SIZE:
                raise InvalidResumeFileError("The expanded DOCX is too large to process")
            names = {member.filename for member in members}
            if "word/document.xml" not in names or "[Content_Types].xml" not in names:
                raise InvalidResumeFileError("The uploaded file is not a Word document")
    except zipfile.BadZipFile as exc:
        raise InvalidResumeFileError("The DOCX is damaged or cannot be opened") from exc


def _append_docx_container(container: Any, blocks: list[str]) -> None:
    for paragraph in container.paragraphs:
        if paragraph.text.strip():
            blocks.append(paragraph.text)
    for table in container.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))


def _validate_size(size: int, max_file_size: int) -> None:
    if max_file_size < 1:
        raise ValueError("max_file_size must be at least 1 byte")
    if size == 0:
        raise InvalidResumeFileError("The uploaded CV is empty")
    if size > max_file_size:
        limit_mb = max_file_size / (1024 * 1024)
        raise ResumeFileTooLargeError(
            f"The uploaded CV exceeds the {limit_mb:g} MB size limit"
        )


def _normalize_text(text: str) -> str:
    normalized_lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    output: list[str] = []
    previous_blank = False
    for line in normalized_lines:
        is_blank = not line
        if is_blank and previous_blank:
            continue
        output.append(line)
        previous_blank = is_blank
    return "\n".join(output).strip()


def _require_text(text: str, resume_format: ResumeFormat) -> None:
    if not text:
        if resume_format is ResumeFormat.PDF:
            raise EmptyResumeTextError(
                "No selectable text was found in the PDF; scanned CVs require OCR"
            )
        raise EmptyResumeTextError("No readable text was found in the DOCX")
