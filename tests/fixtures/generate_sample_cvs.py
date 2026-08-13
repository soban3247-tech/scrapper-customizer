"""Generate fictional CV fixtures for parser integration tests."""

from datetime import UTC, datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfgen.canvas import Canvas
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

FIXTURE_DIR = Path(__file__).resolve().parent / "cvs"

INK = HexColor("#17324D")
MUTED = HexColor("#526271")
DOCX_BLUE = RGBColor(0x2E, 0x74, 0xB5)


def generate_fixtures(output_dir: Path = FIXTURE_DIR) -> tuple[Path, Path]:
    """Create the PDF and DOCX fixtures and return their paths."""

    output_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = output_dir / "synthetic_backend_cv.pdf"
    docx_path = output_dir / "synthetic_data_cv.docx"
    _build_pdf(pdf_path)
    _build_docx(docx_path)
    return pdf_path, docx_path


def _build_pdf(path: Path) -> None:
    styles = getSampleStyleSheet()
    name_style = ParagraphStyle(
        "ResumeName",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=22,
        leading=26,
        textColor=INK,
        alignment=TA_CENTER,
        spaceAfter=4,
    )
    role_style = ParagraphStyle(
        "ResumeRole",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=11,
        leading=14,
        textColor=MUTED,
        alignment=TA_CENTER,
        spaceAfter=16,
    )
    heading_style = ParagraphStyle(
        "ResumeHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        textColor=INK,
        alignment=TA_LEFT,
        spaceBefore=10,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "ResumeBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        textColor=HexColor("#1F2933"),
        spaceAfter=5,
    )

    document = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        rightMargin=0.8 * inch,
        leftMargin=0.8 * inch,
        topMargin=0.7 * inch,
        bottomMargin=0.7 * inch,
        title="Synthetic Backend CV",
        author="Synthetic Fixture Generator",
        subject="Fictional test data only",
    )
    story = [
        Paragraph("Alex Morgan", name_style),
        Paragraph("Senior Backend Developer", role_style),
        Paragraph("SUMMARY", heading_style),
        Paragraph(
            "Senior Backend Developer with 6+ years of professional experience "
            "building reliable REST APIs and cloud services.",
            body_style,
        ),
        Paragraph("SKILLS", heading_style),
        Paragraph(
            "Python, FastAPI, Django, PostgreSQL, Docker, AWS, Git",
            body_style,
        ),
        Paragraph("EXPERIENCE", heading_style),
        Paragraph("Senior Backend Developer | Northwind Labs", body_style),
        Paragraph(
            "Built Python services with FastAPI and PostgreSQL. Deployed "
            "containerized workloads with Docker and AWS.",
            body_style,
        ),
        Spacer(1, 2),
        Paragraph("EDUCATION", heading_style),
        Paragraph(
            "BS Computer Science, Fictional State University, 2018",
            body_style,
        ),
    ]
    document.build(story, canvasmaker=_invariant_canvas)


def _build_docx(path: Path) -> None:
    document = Document()
    document.core_properties.title = "Synthetic Data CV"
    document.core_properties.author = "Synthetic Fixture Generator"
    document.core_properties.subject = "Fictional test data only"
    fixture_timestamp = datetime(2026, 1, 1, tzinfo=UTC)
    document.core_properties.created = fixture_timestamp
    document.core_properties.modified = fixture_timestamp

    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, before, after in (
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = DOCX_BLUE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    # Named override: a simple resume identity block replaces document furniture.
    name = document.add_paragraph()
    name.paragraph_format.space_after = Pt(2)
    name_run = name.add_run("Jordan Lee")
    _format_run(name_run, size=22, bold=True, color=RGBColor(0x17, 0x32, 0x4D))

    role = document.add_paragraph()
    role.paragraph_format.space_after = Pt(14)
    role_run = role.add_run("Data Analyst")
    _format_run(role_run, size=12, color=RGBColor(0x52, 0x62, 0x71))

    document.add_heading("Summary", level=1)
    document.add_paragraph(
        "Data Analyst with 3 years of experience turning operational data into "
        "clear reports and dashboards."
    )

    document.add_heading("Skills", level=1)
    document.add_paragraph("Python, SQL, Excel, Power BI, Tableau")

    document.add_heading("Experience", level=1)
    document.add_paragraph("Data Analyst | Contoso Research")
    document.add_paragraph(
        "Created Power BI and Tableau dashboards, automated Excel reporting, "
        "and analyzed datasets with Python and SQL."
    )

    document.add_heading("Education", level=1)
    document.add_paragraph(
        "Master of Data Science, Fictional State University, 2022"
    )
    document.save(path)


def _format_run(
    run,
    *,
    size: int,
    bold: bool = False,
    color: RGBColor,
) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def _invariant_canvas(filename, **kwargs) -> Canvas:
    """Remove current timestamps from generated PDF metadata."""

    kwargs["invariant"] = 1
    return Canvas(filename, **kwargs)


if __name__ == "__main__":
    generated = generate_fixtures()
    for fixture in generated:
        print(fixture)
